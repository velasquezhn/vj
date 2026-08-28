from __future__ import annotations

from copy import deepcopy
from datetime import date
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "docs" / "manuales" / "01_Manual_General_Villas_Julie.docx"
OUTPUT = ROOT / "docs" / "manuales" / "07_Manual_Maestro_QA_UAT_Villas_Julie.docx"

BLUE = "2563EB"
DARK_BLUE = "17365D"
LIGHT_BLUE = "EAF2F8"
GREEN = "10B981"
LIGHT_GREEN = "E8F5E9"
YELLOW = "FFF4CC"
RED = "FDECEC"
GRAY = "F3F4F6"
WHITE = "FFFFFF"
TEXT = "1F2937"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def set_repeat_table_header(table):
    if table.rows:
        mark_header(table.rows[0])


def style_table(table, header=True):
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if header and r_idx == 0:
                set_cell_shading(cell, DARK_BLUE)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
    if header:
        set_repeat_table_header(table)


def set_run_font(run, size=9.5, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_bullet(doc, text, checked=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(("☐ " if checked else "") + text), 9.5)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(text), 9.5)
    return p


def add_label_paragraph(doc, label, value=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(label + ": "), 9.5, True, DARK_BLUE)
    set_run_font(p.add_run(value), 9.5)
    return p


def add_callout(doc, title, text, color=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    set_run_font(p.add_run(title + "\n"), 10, True, DARK_BLUE)
    set_run_font(p.add_run(text), 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def clear_body_keep_section(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


MODULE_CONTEXT = {
    "Preparación": "documentación y ambiente QA",
    "Autenticación": "pantalla Inicio de sesión o menú de perfil",
    "Dashboard": "menú Dashboard",
    "Reservas": "menú Reservas",
    "Crear reservación": "menú Crear Reservación",
    "Calendario": "menú Calendario o Calendario Mejorado",
    "Usuarios": "menú Usuarios",
    "Cabañas": "menú Cabañas",
    "Tipos de menú": "menú Tipos de Menú",
    "Actividades": "menú Actividades",
    "Reportes": "menú Reportes",
    "Administradores": "menú Administradores",
    "Configuración": "menú Configuración",
    "Copias": "menú Copias de seguridad",
    "Mensajes pendientes": "menú Mensajes pendientes",
    "Auditoría": "menú Auditoría",
    "Estados de conversación": "ruta protegida /conversation-states",
    "WhatsApp huésped": "chat del número oficial de Villas Julie",
    "WhatsApp administrador": "chat del número oficial desde un teléfono administrador",
    "Webhook Meta": "ambiente QA y herramientas técnicas autorizadas",
    "Seguridad": "panel, API QA y herramientas técnicas autorizadas",
    "Interfaz": "panel administrativo en el dispositivo indicado",
    "Regresión": "ambiente QA completo",
    "Producción": "producción, únicamente con pruebas no destructivas",
}

ROLE_PRE = {
    "Superadministrador": ["Cuenta QA con rol superadministrador activa.", "Contraseña QA conocida por el probador y no escrita en este manual."],
    "Administrador": ["Cuenta QA con rol admin activa.", "La cuenta no debe poseer permisos de superadministrador."],
    "Administrador WhatsApp": ["Número QA agregado y activo como administrador de WhatsApp.", "El teléfono debe haber enviado /admin para abrir la ventana de atención."],
    "Huésped": ["Número de teléfono QA autorizado por Meta.", "No usar el número personal de un cliente real."],
    "Superadministrador + técnico QA": ["Cuenta QA superadministrador.", "Apoyo técnico con acceso al ambiente QA, sin revelar secretos."],
    "Todos los roles": ["Las cuentas y teléfonos QA requeridos deben estar preparados."],
}

TEST_DATA = {
    "Autenticación": [("Usuario", "[USUARIO QA DEL ROL]"), ("Contraseña", "[CONTRASEÑA QA — NO ESCRIBIR EN EL MANUAL]")],
    "Reservas": [("Huésped", "QA Huésped 001"), ("Teléfono", "[TELÉFONO QA AUTORIZADO]"), ("Cabaña", "Tortuga o una cabaña libre del ambiente QA"), ("Fechas", "Entrada dentro de 90 días; salida dos días después")],
    "Crear reservación": [("Huésped", "QA Huésped 001"), ("Teléfono", "[TELÉFONO QA AUTORIZADO]"), ("Personas", "2"), ("Precio", "Calculado por el sistema")],
    "Usuarios": [("Nombre", "QA Huésped 001"), ("Teléfono", "[TELÉFONO QA ÚNICO]"), ("Rol", "guest")],
    "Cabañas": [("Nombre", "QA Cabaña 001"), ("Capacidad", "4"), ("Precio", "2500"), ("Foto", "JPG QA menor de 5 MiB")],
    "Tipos de menú": [("Nombre", "QA Tipo 001"), ("Capacidad", "4"), ("Precio", "2500"), ("Fotos", "Dos URL HTTPS QA, una por línea")],
    "Actividades": [("Clave", "qa_actividad_001"), ("Nombre", "Actividad QA 001"), ("Categoría", "Naturaleza"), ("Multimedia", "JSON válido con URL HTTPS QA")],
    "Administradores": [("Usuario", "qa_admin_001"), ("Correo", "qa_admin_001@example.invalid"), ("Contraseña", "[CONTRASEÑA QA SEGURA DEFINIDA POR EL PROPIETARIO]"), ("Rol", "admin")],
    "Configuración": [("Administrador", "QA Administrador WhatsApp"), ("Número", "[NÚMERO QA CON 504]"), ("Anticipo", "50"), ("Cuenta", "Banco QA / Cuenta QA / 000000 — SOLO STAGING")],
    "WhatsApp huésped": [("Número huésped", "[TELÉFONO QA AUTORIZADO]"), ("Mensaje inicial", "hola")],
    "WhatsApp administrador": [("Número administrador", "[TELÉFONO QA ADMINISTRADOR]"), ("Código", "[VJ-XXXXXX CREADO DURANTE LA PRUEBA]")],
}


tests = []


def add(group, role, module, title, action, expected, criticality="Media", test_type="Funcional", data=None, pre=None, verify=None, failure=None, evidence=None, technical=False):
    test_id = len(tests) + 1
    requirements = deepcopy(ROLE_PRE.get(role, []))
    if pre:
        requirements.extend(pre)
    if module not in ("Preparación", "Producción"):
        requirements.append("Usar el ambiente QA; no modificar datos de producción.")
    base_data = deepcopy(TEST_DATA.get(module, []))
    if data:
        base_data.extend(data)
    context = MODULE_CONTEXT[module]
    steps = [
        "Abra el ambiente indicado en Preparación y confirme que corresponde a QA.",
    ]
    if role in ("Superadministrador", "Administrador", "Superadministrador + técnico QA"):
        steps.append(f"Inicie sesión con el rol {role.replace(' + técnico QA', '')}.")
    elif role == "Administrador WhatsApp":
        steps.append("Abra WhatsApp en el teléfono administrador QA y entre al chat oficial de Villas Julie.")
    elif role == "Huésped":
        steps.append("Abra WhatsApp en el teléfono huésped QA y entre al chat oficial de Villas Julie.")
    elif role == "Todos los roles":
        steps.append("Prepare las sesiones de cada rol indicadas por la prueba.")
    steps.append(f"Ubíquese en {context}.")
    steps.extend([part.strip() for part in action.split("|") if part.strip()])
    steps.append("Espere la respuesta completa; no presione el botón nuevamente mientras esté procesando.")
    if module not in ("Preparación", "Webhook Meta", "Seguridad", "Producción"):
        steps.append("Actualice la pantalla o vuelva a abrir el menú y confirme que el resultado permanece.")
    checks = verify or [
        expected,
        "El mensaje visible es claro y no muestra trazas, SQL, tokens ni errores internos.",
        "El resultado permanece al actualizar o volver a consultar.",
    ]
    failures = failure or [
        "La acción no responde, devuelve un error interno o queda cargando indefinidamente.",
        "El resultado no coincide con lo esperado o desaparece al actualizar.",
        "Se modifica información distinta a la seleccionada.",
    ]
    tests.append({
        "id": test_id,
        "group": group,
        "role": role,
        "module": module,
        "title": title,
        "criticality": criticality,
        "type": test_type,
        "objective": f"Comprobar que {title.lower()}.",
        "requirements": requirements,
        "data": base_data or [("Dato", "No aplica; seguir el procedimiento")],
        "steps": steps,
        "expected": expected,
        "verify": checks,
        "failure": failures,
        "evidence": evidence or "Captura de la pantalla o conversación antes y después, con datos sensibles ocultos.",
        "technical": technical,
    })


def add_many(group, role, module, rows, criticality="Media", test_type="Funcional"):
    for row in rows:
        title, action, expected, *extras = row
        kwargs = extras[0] if extras else {}
        add(group, role, module, title, action, expected, criticality, test_type, **kwargs)


# BLOQUE 1 — Preparación
add_many("BLOQUE 1 — Preparación", "Superadministrador + técnico QA", "Preparación", [
    ("identificar correctamente las URL de QA y producción", "Anote la URL QA y la URL de producción|Abra cada una y compruebe el nombre del ambiente", "Las URL quedan diferenciadas y ninguna prueba destructiva se ejecutará en producción."),
    ("crear una copia de la base antes de las pruebas", "Pulse Crear ahora en Copias de seguridad del ambiente QA|Descargue el archivo y registre nombre, fecha y tamaño", "Existe una copia descargable, no vacía y anterior a cualquier cambio QA."),
    ("preparar una base QA separada", "Confirme con el responsable técnico que DB_PATH apunta a una base QA|Verifique que no contiene clientes ni pagos reales", "El ambiente QA usa almacenamiento separado de producción."),
    ("preparar las cuentas de los cuatro roles", "Complete la tabla de credenciales sin escribir contraseñas|Confirme superadmin, admin, administrador WhatsApp y huésped", "Todos los roles reales tienen un medio de prueba disponible."),
    ("preparar datos seguros y fechas libres", "Seleccione dos fechas futuras libres y una cabaña QA|Reserve nombres con prefijo QA", "Los datos de prueba están identificados y no chocan con reservas reales."),
    ("preparar archivos de comprobante", "Prepare JPG, PNG y PDF válidos menores de 5 MiB|Prepare un archivo inválido y uno mayor de 5 MiB", "Los cinco archivos de prueba están disponibles y no contienen datos bancarios reales."),
    ("preparar navegadores y teléfonos", "Prepare Chrome o Edge de escritorio, una ventana privada y un teléfono Android o iPhone|Registre versiones", "Hay dispositivos suficientes para pruebas de sesión, permisos y WhatsApp."),
    ("definir el mecanismo de limpieza QA", "Liste los registros QA que podrán eliminarse|Asegure que solo el responsable QA hará la limpieza", "Existe una lista de limpieza y una copia previa para recuperación."),
], "Alta", "Preparación")

# BLOQUE 2 — Autenticación y sesiones
add_many("BLOQUE 2 — Autenticación y sesiones", "Superadministrador", "Autenticación", [
    ("iniciar sesión con credenciales correctas", "Escriba Usuario y Contraseña válidos|Pulse Entrar", "El panel abre Dashboard y muestra el menú del superadministrador."),
    ("rechazar una contraseña incorrecta", "Escriba un usuario válido y una contraseña incorrecta|Pulse Entrar", "Permanece en Inicio de sesión y muestra Credenciales inválidas sin revelar cuál campo falló."),
    ("rechazar un usuario inexistente", "Escriba qa_usuario_inexistente y cualquier contraseña QA|Pulse Entrar", "Permanece en Inicio de sesión y muestra Credenciales inválidas."),
    ("validar campos de acceso vacíos", "Deje Usuario y Contraseña vacíos|Pulse Entrar", "El formulario no inicia sesión y señala que los campos son obligatorios."),
    ("aplicar límite de intentos fallidos", "En QA, repita intentos incorrectos hasta alcanzar el límite configurado|Registre la respuesta", "El sistema responde Demasiados intentos o HTTP 429 sin bloquear permanentemente la cuenta.", {"technical": True}),
    ("cerrar sesión desde el menú de perfil", "Abra el avatar|Pulse Cerrar sesión", "La sesión termina y vuelve a Inicio de sesión."),
    ("impedir volver al panel después de cerrar sesión", "Después de cerrar sesión pulse Atrás en el navegador|Abra directamente /backups", "El sistema redirige a /login y no muestra información protegida."),
    ("proteger una URL sin sesión", "Abra una ventana privada|Escriba directamente /reservations", "La ruta redirige a /login."),
    ("mantener sesión al actualizar", "Inicie sesión|Actualice Dashboard", "La sesión válida continúa y el rol se conserva."),
    ("sincronizar cierre de sesión entre pestañas", "Abra el panel en dos pestañas|Cierre sesión en una|Actualice la otra", "La segunda pestaña deja de acceder a datos protegidos."),
    ("cambiar la contraseña propia", "Abra el avatar|Pulse Cambiar contraseña|Complete Contraseña actual, Contraseña nueva y Confirmar contraseña nueva|Guarde", "La contraseña cambia, la sesión anterior se invalida y se solicita iniciar sesión nuevamente."),
    ("rechazar cambio con contraseña actual incorrecta", "Abra Cambiar contraseña|Escriba una contraseña actual incorrecta y una nueva válida|Guarde", "Muestra que la contraseña actual es incorrecta y no cambia la credencial."),
    ("rechazar una contraseña nueva débil", "Escriba una nueva contraseña menor de 10 caracteres o sin letras/números|Guarde", "Muestra la regla de al menos 10 caracteres, letras y números."),
    ("forzar cambio en una cuenta nueva", "Inicie sesión con un administrador QA recién creado", "El sistema dirige a Cambiar contraseña antes de permitir el resto del panel."),
], "Alta", "Autenticación / Seguridad")

# BLOQUE 3 — Dashboard
add_many("BLOQUE 3 — Dashboard", "Administrador", "Dashboard", [
    ("mostrar métricas globales", "Abra Dashboard|Pulse Actualizar", "Se muestran Reservas registradas, Valor de reservas confirmadas, Huéspedes registrados y Ocupación últimos 30 días."),
    ("mostrar tareas pendientes", "Cree previamente una solicitud en cada estado pendiente|Abra Dashboard", "Los contadores de autorización, pago vencido, comprobante y mensajes reflejan los datos QA."),
    ("abrir el listado desde Ver todas", "Pulse Ver todas en reservas recientes", "Se abre Reservas y el listado coincide con el dashboard."),
    ("actualizar sin duplicar métricas", "Pulse Actualizar dos veces con pausa|Compare valores", "Los valores no se duplican y permanecen consistentes."),
    ("manejar un dashboard sin reservas", "Use una base QA vacía|Abra Dashboard", "Muestra ceros y estados vacíos sin error técnico."),
    ("conservar precisión después de confirmar una reserva", "Registre los valores|Confirme una reserva QA|Pulse Actualizar", "El valor confirmado y los contadores cambian exactamente una vez."),
], "Alta", "Funcional / Datos")

# BLOQUE 4 — Reservas
add_many("BLOQUE 4 — Reservas", "Administrador", "Reservas", [
    ("listar reservas", "Abra Reservas", "La tabla muestra código/ID, huésped, teléfono, cabaña, fechas, personas, total, estado y acciones."),
    ("buscar por nombre", "Escriba QA Huésped 001 en Buscar por usuario, teléfono, cabaña o ID", "Solo aparecen coincidencias por nombre."),
    ("buscar por teléfono", "Escriba el teléfono QA completo", "Aparece la reserva asociada al teléfono."),
    ("buscar por código o ID", "Escriba el código VJ o ID QA", "Aparece la reserva exacta."),
    ("filtrar por estado", "Abra Filtros|Seleccione pendiente_autorizacion", "Solo aparecen solicitudes en espera de autorización."),
    ("filtrar por cabaña y fechas", "Seleccione Cabaña, Desde y Hasta|Aplique", "La tabla muestra únicamente reservas dentro de los filtros."),
    ("limpiar todos los filtros", "Aplique búsqueda y filtros|Pulse Limpiar Todo", "Se restablece el listado completo."),
    ("abrir y cancelar Nueva Reserva", "Pulse Nueva Reserva|Sin guardar pulse Cancelar", "El diálogo se cierra y no se crea ningún registro."),
    ("crear una reserva válida desde Reservas", "Pulse Nueva Reserva|Complete ID de Usuario, ID de Cabaña, fechas, personas y Precio Total|Pulse Crear", "Se crea una reserva QA en estado pendiente_autorizacion y aparece en la tabla."),
    ("rechazar una reserva sin campos obligatorios", "Pulse Nueva Reserva|Deje campos obligatorios vacíos|Pulse Crear", "El formulario señala los campos requeridos y no crea registro."),
    ("rechazar una fecha de salida anterior", "Use salida igual o anterior a la entrada|Pulse Crear", "La reserva no se crea y aparece un mensaje de fechas inválidas."),
    ("rechazar cero personas", "Escriba 0 en Número de Personas|Pulse Crear", "La reserva no se crea y se informa un valor inválido."),
    ("rechazar precio negativo", "Escriba -1 en Precio Total|Pulse Crear", "La reserva no se crea o el campo impide el valor negativo."),
    ("impedir fechas superpuestas", "Cree una reserva QA válida|Intente otra para la misma cabaña con fechas que se cruzan", "La segunda operación se rechaza por conflicto y la primera permanece intacta."),
    ("permitir fechas contiguas sin cruce", "Use una entrada igual a la salida de otra reserva|Guarde", "La reserva se crea si la política de noches no produce superposición."),
    ("editar una reserva válida", "Pulse Editar reserva|Cambie personas o precio|Pulse Actualizar", "La tabla muestra los nuevos valores y conserva el código."),
    ("cancelar una edición", "Pulse Editar reserva|Cambie un dato|Pulse Cancelar", "No se guarda el cambio."),
    ("impedir conflicto al editar fechas", "Edite fechas para cruzarlas con otra reserva activa|Pulse Actualizar", "La edición se rechaza y las fechas originales permanecen."),
    ("autorizar el pago de una solicitud", "En Espera autorización pulse Autorizar pago y notificar al huésped", "Cambia a Pago autorizado, calcula 50 %, fija vencimiento de 24 horas y notifica o encola el aviso."),
    ("bloquear autorización sin cuentas de pago", "En QA deje Cuentas bancarias vacías|Intente Autorizar pago", "La acción se bloquea con una explicación de configuración incompleta."),
    ("impedir comprobante antes de autorización", "Desde el teléfono QA envíe una foto cuando la reserva esté pendiente_autorizacion", "El sistema rechaza el archivo como comprobante y mantiene el estado."),
    ("abrir un comprobante autorizado", "Después de enviarlo, pulse Abrir comprobante", "Se abre el archivo correcto y no otro comprobante."),
    ("confirmar después de revisar el comprobante", "Pulse Confirmar después de revisar el comprobante", "La reserva cambia a confirmada, registra revisión y notifica o encola al huésped."),
    ("impedir confirmar sin comprobante", "Intente confirmar una reserva sin archivo", "La acción se rechaza y la reserva no queda confirmada."),
    ("rechazar una solicitud con motivo", "Pulse Rechazar y notificar|Escriba un motivo claro|Confirme", "La reserva cambia a rechazada, conserva el motivo y notifica o encola al huésped."),
    ("rechazar un motivo demasiado corto", "Pulse Rechazar y escriba menos de tres caracteres", "No ejecuta el rechazo y solicita un motivo válido."),
    ("evitar doble confirmación", "Pulse Confirmar y vuelva a intentarlo inmediatamente", "Solo se registra una transición; la segunda acción no duplica mensajes ni cambios."),
    ("eliminar una reserva QA con confirmación", "Seleccione una reserva creada solo para QA|Pulse Eliminar reserva|Confirme", "El registro QA desaparece y no afecta otras reservas."),
    ("cancelar la eliminación de una reserva", "Pulse Eliminar reserva|Cancele en la confirmación", "La reserva permanece sin cambios."),
    ("exportar CSV visible", "Aplique filtros|Pulse Exportar CSV", "Se descarga un CSV con encabezados y únicamente las filas filtradas."),
], "Alta", "CRUD / Flujo / Validación")

# BLOQUE 5 — Formulario avanzado de reserva
add_many("BLOQUE 5 — Crear reservación", "Administrador", "Crear reservación", [
    ("crear una reserva guiada con usuario existente", "Seleccione un huésped|Seleccione Cabaña, personas y fechas|Revise Precio Total|Pulse Crear Reserva", "La reserva aparece en Reservas y Calendario."),
    ("crear un huésped desde el formulario", "Pulse Crear nuevo usuario|Complete Nombre Completo y Teléfono|Guarde|Continúe la reserva", "El huésped se crea una vez y queda seleccionado."),
    ("rechazar un teléfono inválido", "Pulse Crear nuevo usuario|Escriba letras o un número incompleto|Guarde", "El usuario no se crea y el teléfono se marca inválido."),
    ("evitar teléfono duplicado", "Intente crear un usuario con un teléfono existente", "No duplica al huésped; permite seleccionar el registro existente o muestra error claro."),
    ("calcular precio por fechas", "Seleccione cabaña y fechas válidas", "Precio Total se actualiza según noches y tarifa configurada."),
    ("mostrar fechas ocupadas", "Seleccione una cabaña con reserva QA|Abra los selectores de fecha", "Las fechas ocupadas se identifican o la creación se rechaza al intentar usarlas."),
    ("cancelar la creación guiada", "Complete parcialmente el formulario|Pulse Cancelar", "Vuelve sin crear reserva ni usuario adicional."),
], "Alta", "Funcional / Validación")

# BLOQUE 6 — Calendario
add_many("BLOQUE 6 — Calendarios", "Administrador", "Calendario", [
    ("mostrar ocupación mensual", "Abra Calendario", "Cada cabaña y día muestra Libre o el estado de la reserva correspondiente."),
    ("navegar al mes anterior y siguiente", "Pulse Anterior|Pulse Siguiente", "El encabezado y la ocupación cambian al mes correcto."),
    ("volver al mes actual", "En Calendario Mejorado navegue a otro mes|Pulse Hoy", "Vuelve al mes y fecha actuales."),
    ("filtrar por tipo y capacidad", "Seleccione Tipo de Cabaña y Capacidad mínima", "Solo permanecen las cabañas que cumplen ambos filtros."),
    ("filtrar por estado", "Seleccione Confirmado o Pendiente", "La visualización refleja únicamente el estado elegido."),
    ("alternar cuadrícula y tabla", "Use Vista de cuadrícula|Use Vista de tabla", "Ambas vistas representan las mismas reservas."),
    ("abrir detalles de una reserva", "Pulse una celda ocupada|Revise Detalles de la Reserva|Pulse Gestionar Reserva", "Muestra datos correctos y abre la reserva correspondiente."),
    ("limpiar filtros", "Aplique varios filtros|Pulse Limpiar Filtros", "Vuelve a mostrar todas las cabañas."),
], "Media", "Funcional / Interfaz")

# BLOQUE 7 — Usuarios
add_many("BLOQUE 7 — Usuarios huésped", "Administrador", "Usuarios", [
    ("listar huéspedes", "Abra Usuarios", "Se muestran nombre, teléfono, rol, estado y reservas."),
    ("buscar un huésped", "Escriba nombre o teléfono QA|Pulse Buscar", "Aparece únicamente el huésped coincidente."),
    ("filtrar por estado y rol", "Seleccione Estado y Rol", "El listado respeta ambos filtros."),
    ("crear un huésped válido", "Pulse Nuevo Usuario|Complete Nombre Completo, Teléfono y Rol guest|Guarde", "El huésped aparece y persiste después de actualizar."),
    ("rechazar teléfono duplicado", "Intente crear otro huésped con el mismo teléfono", "No se crea un duplicado."),
    ("editar nombre y estado", "Abra un huésped QA|Cambie Nombre Completo o Estado|Pulse Actualizar", "Los cambios persisten y no cambian su teléfono por accidente."),
    ("cancelar edición de huésped", "Abra editar|Cambie un campo|Pulse Cancelar", "El registro conserva los datos anteriores."),
    ("actualizar estados desde reservas", "Pulse Actualizar Estados", "El sistema completa la operación una sola vez y muestra confirmación sin alterar teléfonos ni nombres."),
], "Media", "CRUD / Validación")

# BLOQUE 8 — Cabañas
add_many("BLOQUE 8 — Cabañas", "Administrador", "Cabañas", [
    ("listar las 13 cabañas reales", "Abra Cabañas|Seleccione todas", "Se muestran las 13 cabañas configuradas, sin duplicados."),
    ("filtrar por tipo", "Use filtro de cabañas y seleccione tortuga, delfin u otro disponible", "Solo aparecen cabañas del tipo seleccionado."),
    ("crear una cabaña QA", "Pulse Nueva Cabaña|Complete Nombre, Capacidad, Precio por noche, Descripción y Disponible para reservas|Guarde", "La cabaña QA aparece en el panel."),
    ("crear una cabaña con foto JPG", "Adjunte JPG QA menor de 5 MiB|Guarde", "La imagen se carga y aparece en la tarjeta correcta."),
    ("rechazar imagen demasiado grande", "Adjunte una imagen mayor de 5 MiB|Guarde", "La carga se rechaza con mensaje comprensible."),
    ("rechazar campos vacíos", "Pulse Nueva Cabaña|Deje Nombre o Capacidad vacío|Guarde", "No se crea el registro."),
    ("rechazar capacidad o precio inválido", "Use capacidad 0 o precio negativo|Guarde", "No se crea o actualiza con valores inválidos."),
    ("editar una cabaña", "Pulse Editar|Cambie descripción, precio o disponibilidad|Pulse Actualizar", "La información persiste y WhatsApp refleja el cambio cuando corresponde."),
    ("cancelar edición de cabaña", "Pulse Editar|Cambie un dato|Pulse Cancelar", "No se guarda el cambio."),
    ("impedir eliminar una cabaña relacionada", "En QA seleccione una cabaña con reserva|Pulse Eliminar|Confirme", "La eliminación se rechaza o conserva integridad; ninguna reserva queda huérfana."),
    ("eliminar una cabaña QA sin relaciones", "Cree una cabaña QA sin reservas|Pulse Eliminar|Confirme", "Solo la cabaña QA se elimina."),
    ("cancelar eliminación de cabaña", "Pulse Eliminar|Cancele", "La cabaña permanece."),
], "Alta", "CRUD / Archivos / Integridad")

# BLOQUE 9 — Tipos de menú
add_many("BLOQUE 9 — Tipos de menú WhatsApp", "Administrador", "Tipos de menú", [
    ("listar tipos de alojamiento", "Abra Tipos de Menú", "Se muestran nombre, capacidad, habitaciones, baños, precio, imágenes y estado."),
    ("abrir vista previa de WhatsApp", "Pulse Vista Previa del Menú", "La vista usa el contenido activo y no muestra JSON técnico."),
    ("editar texto, capacidad y precio", "Pulse Editar|Cambie campos QA|Pulse Guardar", "Los cambios persisten y aparecen en la vista previa."),
    ("editar varias fotos HTTPS", "En URLs de Fotos agregue dos direcciones HTTPS, una por línea|Guarde", "La cantidad de imágenes aumenta y ambas son accesibles."),
    ("rechazar una URL inválida", "Agregue texto que no sea URL HTTPS|Guarde", "El sistema rechaza el valor o la vista previa indica claramente el error."),
    ("activar y desactivar un tipo", "Cambie el interruptor Estado|Confirme el menú", "El tipo inactivo deja de aparecer al huésped y vuelve al activarlo."),
    ("cancelar edición del tipo", "Pulse Editar|Cambie datos|Pulse Cancelar", "Los datos originales permanecen."),
    ("verificar galería para cada tipo", "Abra Tortuga, Delfín y los demás tipos activos|Revise fotos", "Cada tipo envía o muestra sus propias imágenes; no reutiliza únicamente las de Tortuga."),
], "Alta", "CRUD / WhatsApp / Multimedia")

# BLOQUE 10 — Actividades
add_many("BLOQUE 10 — Actividades", "Administrador", "Actividades", [
    ("listar actividades", "Abra Actividades", "La tabla muestra foto, nombre, categoría, duración, precio, capacidad y acciones."),
    ("crear una actividad válida", "Pulse Nueva Actividad|Complete los campos de prueba|Pulse Crear Actividad", "La actividad aparece en la tabla y en la vista previa si está activa."),
    ("rechazar clave duplicada", "Cree otra actividad con qa_actividad_001", "No se crea duplicado y aparece un mensaje claro."),
    ("rechazar JSON inválido", "Escriba texto no JSON en URLs de Fotos|Pulse Crear Actividad", "No se guarda y se identifica el campo inválido."),
    ("editar una actividad", "Pulse Editar|Cambie Nombre, Descripción o Multimedia|Pulse Actualizar", "Los cambios persisten."),
    ("cancelar edición de actividad", "Pulse Editar|Cambie un dato|Pulse Cancelar", "No se guarda el cambio."),
    ("abrir galería de fotos", "Pulse Ver fotos", "Abre la galería de la actividad elegida y permite cerrar."),
    ("abrir vista previa individual", "Pulse Vista previa del menú", "Muestra el mensaje y botones que recibirá el huésped."),
    ("abrir vista previa general del bot", "Pulse Vista Previa Bot", "Muestra únicamente actividades activas e incluidas en menú."),
    ("activar y desactivar actividad", "Cambie el estado activo|Vuelva a abrir la vista previa", "La actividad desaparece y reaparece correctamente."),
    ("eliminar una actividad QA", "Seleccione una actividad QA sin relaciones|Pulse Eliminar|Confirme", "La actividad QA desaparece sin afectar otras."),
], "Media", "CRUD / JSON / WhatsApp")

# BLOQUE 11 — Reportes
add_many("BLOQUE 11 — Reportes", "Administrador", "Reportes", [
    ("consultar reportes por período", "Seleccione Desde y Hasta|Pulse Aplicar", "Muestra ingresos, ocupación, usuarios y tendencias del período."),
    ("rechazar un período invertido", "Use Hasta anterior a Desde|Pulse Aplicar", "No consulta y muestra una validación de fechas."),
    ("mostrar período sin datos", "Elija fechas QA sin reservas|Pulse Aplicar", "Muestra ceros o Sin datos sin fallar."),
    ("calcular valor confirmado", "Use un período con una reserva confirmada conocida|Compare el total", "El valor coincide con el total contractual de reservas confirmadas."),
    ("calcular ocupación por noches", "Use una reserva QA de dos noches|Compare ocupación", "La ocupación cuenta noches reservadas, no días duplicados."),
    ("exportar Excel", "Aplique período|Pulse Descargar Excel", "Se descarga un XLSX que abre y contiene las reservas del período."),
    ("mantener filtros al actualizar", "Aplique fechas|Actualice la consulta", "Los resultados corresponden al mismo intervalo."),
], "Alta", "Reportes / Exportación")

# BLOQUE 12 — Administradores del panel
add_many("BLOQUE 12 — Superadministrador", "Superadministrador", "Administradores", [
    ("listar administradores", "Abra Administradores", "Se muestran username, email, nombre, rol, estado y fecha."),
    ("crear un administrador", "Pulse Nuevo Administrador|Complete datos QA con rol Admin|Guarde", "La cuenta se crea activa y exige cambio de contraseña en el primer acceso."),
    ("crear un superadministrador QA", "Pulse Nuevo Administrador|Seleccione Superadministrador|Guarde", "La cuenta se crea con rol superadmin y menú superior después del cambio obligatorio."),
    ("rechazar username o email duplicado", "Intente crear una cuenta con username o email existente", "No se crea y aparece el mensaje de duplicado."),
    ("rechazar contraseña débil al crear", "Use una contraseña menor de 10 caracteres o sin letras/números|Guarde", "No se crea la cuenta."),
    ("editar datos y rol", "Pulse Editar|Cambie Nombre Completo o Rol|Guarde", "Los cambios persisten y el menú cambia según el rol al verificar sesión."),
    ("cancelar edición administrativa", "Pulse Editar|Cambie datos|Pulse Cancelar", "No se guarda."),
    ("cambiar contraseña de otra cuenta", "Pulse Cambiar contraseña|Escriba y confirme una nueva válida|Guarde", "La cuenta objetivo debe cambiarla al iniciar y sus sesiones anteriores quedan inválidas."),
    ("activar y desactivar otra cuenta", "Pulse Desactivar|Compruebe que no puede iniciar|Vuelva a Activar", "El acceso se bloquea y se recupera de acuerdo con el estado."),
    ("impedir desactivarse a sí mismo", "En su propia fila intente Desactivar", "La operación se bloquea."),
    ("impedir eliminarse a sí mismo", "En su propia fila intente Eliminar", "La operación se bloquea."),
    ("proteger el último superadministrador", "En QA deje un solo superadmin activo|Intente degradarlo, desactivarlo y eliminarlo", "Las tres operaciones se rechazan y el panel conserva un superadministrador activo."),
    ("eliminar lógicamente una cuenta QA", "Seleccione otro administrador QA|Pulse Eliminar|Confirme", "La cuenta queda inactiva y no puede iniciar sesión; el historial se conserva."),
], "Alta", "CRUD / Roles / Seguridad")

# BLOQUE 13 — Configuración
add_many("BLOQUE 13 — Configuración sensible", "Superadministrador", "Configuración", [
    ("listar administradores de WhatsApp", "Abra Configuración", "Se muestran únicamente números autorizados activos/inactivos con nombre."),
    ("agregar número de administrador válido", "Pulse Agregar número|Complete Nombre y Número con código de país|Pulse Guardar", "El número se guarda y el sistema intenta enviar la prueba."),
    ("rechazar teléfono inválido", "Agregue letras o un número sin formato 504XXXXXXXX|Pulse Guardar", "No guarda y muestra formato inválido."),
    ("rechazar número duplicado", "Agregue un número ya registrado", "No duplica y muestra que ya existe."),
    ("editar nombre o estado", "Pulse Editar|Cambie Nombre o estado|Guarde", "Los cambios persisten."),
    ("enviar prueba al administrador", "Pulse Enviar prueba", "El panel informa enviado o explica que el teléfono debe escribir /admin; nunca afirma éxito falso."),
    ("desactivar recepción administrativa", "Edite un número y desactívelo|Cree una solicitud QA", "El número desactivado no recibe nuevas revisiones."),
    ("eliminar un número QA", "Pulse Eliminar sobre un número creado para QA|Confirme", "El número desaparece y deja de estar autorizado."),
    ("cancelar eliminación de número", "Pulse Eliminar|Cancele", "El número permanece."),
    ("guardar anticipo del 50 por ciento", "En Porcentaje de anticipo escriba 50|Pulse Guardar pagos", "El valor persiste y una autorización nueva calcula exactamente 50 %."),
    ("rechazar porcentaje fuera de rango", "Escriba 0, 101 o texto|Pulse Guardar pagos", "No guarda y muestra valor inválido."),
    ("guardar cuentas bancarias de QA", "En Cuentas bancarias escriba una cuenta por línea marcada QA|Pulse Guardar pagos", "Las cuentas persisten y aparecen en Pago autorizado del ambiente QA."),
    ("guardar instrucciones adicionales", "Escriba una nota QA|Pulse Guardar pagos", "La nota persiste y aparece junto a las instrucciones de pago."),
    ("bloquear Configuración a un admin normal", "Cierre sesión|Entre como Admin|Escriba directamente /settings", "El menú no aparece, la URL redirige y la API responde 403."),
], "Crítica", "Configuración / Permisos / WhatsApp")

# BLOQUE 14 — Copias
add_many("BLOQUE 14 — Copias de seguridad", "Superadministrador", "Copias", [
    ("consultar estado de copias", "Abra Copias de seguridad|Pulse Actualizar", "Muestra servicio, cantidad, última copia y estado Dropbox."),
    ("crear una copia manual", "Anote la cantidad|Pulse Crear ahora", "Aparece una nueva copia con fecha actual y tamaño mayor de 0.00 MB."),
    ("evitar doble creación accidental", "Pulse Crear ahora y trate de pulsarlo nuevamente mientras dice Creando", "El botón permanece deshabilitado y no inicia dos copias simultáneas."),
    ("descargar una copia", "Pulse Descargar en la copia más reciente", "Se descarga un SQLite o comprimido con nombre y tamaño coherentes."),
    ("validar integridad técnica de la copia", "Entregue el archivo al técnico QA|Abra una copia aislada y liste sus tablas", "La copia contiene Users, Admins, Cabins, Reservations y las tablas operativas.", {"technical": True}),
    ("proteger descarga sin sesión", "Cierre sesión|Intente repetir la URL de descarga", "La descarga se rechaza con 401/403."),
    ("bloquear copias a un admin normal", "Entre como Admin|Escriba /backups", "El menú no aparece, la URL redirige y la API rechaza."),
    ("mantener restauración fuera de línea", "En QA solicite técnicamente POST /admin/backup/restore sin detener servicio", "Responde OFFLINE_RESTORE_REQUIRED y no modifica la base.", {"technical": True}),
], "Crítica", "Backup / Seguridad / Persistencia")

# BLOQUE 15 — Cola, auditoría y estados
add_many("BLOQUE 15 — Operación superior", "Superadministrador", "Mensajes pendientes", [
    ("listar mensajes de WhatsApp", "Abra Mensajes pendientes|Pulse Actualizar", "La tabla muestra destinatario enmascarado, tipo, estado, intentos, próximo intento y error."),
    ("filtrar mensajes por estado", "Seleccione Estado pending, failed o sent", "Solo aparecen filas del estado elegido."),
    ("reintentar un mensaje fallido", "En QA seleccione un mensaje reintentable|Pulse Reintentar", "Aumenta intentos, cambia el estado y registra la respuesta de Meta."),
    ("impedir reintentar un mensaje enviado", "Localice un mensaje sent|Revise acciones", "No permite duplicar una notificación ya enviada."),
    ("bloquear la cola a un admin normal", "Entre como Admin|Escriba /notifications", "No aparece el menú y el acceso se rechaza."),
], "Alta", "Notificaciones / Reintentos / Permisos")
add_many("BLOQUE 15 — Operación superior", "Superadministrador", "Auditoría", [
    ("listar eventos administrativos", "Abra Auditoría|Pulse Actualizar", "Muestra fecha, usuario, rol, método, ruta y código de resultado."),
    ("registrar una modificación", "Cambie un registro QA|Vuelva a Auditoría", "Aparece un evento con usuario y resultado correctos."),
    ("registrar un intento rechazado", "Provoque una validación inválida no destructiva|Vuelva a Auditoría", "Aparece el método/ruta y un código de error sin contraseña ni token."),
    ("bloquear auditoría a un admin normal", "Entre como Admin|Escriba /audit-logs", "No aparece el menú y el acceso se rechaza."),
], "Alta", "Auditoría / Seguridad")
add_many("BLOQUE 15 — Operación superior", "Superadministrador", "Estados de conversación", [
    ("listar estados persistentes", "Abra /conversation-states", "Se muestran Número de Usuario, Estado y Datos JSON."),
    ("crear y editar un estado QA", "Pulse Nuevo Estado|Complete Número, Estado y JSON válido|Guarde|Edite y vuelva a guardar", "El estado persiste con JSON válido."),
    ("rechazar JSON inválido", "En Datos escriba JSON incompleto|Guarde", "No guarda o muestra un error claro."),
    ("eliminar un estado QA", "Seleccione únicamente el estado QA|Pulse Eliminar|Confirme", "El estado QA desaparece."),
    ("bloquear estados a un admin normal", "Entre como Admin|Escriba /conversation-states", "El acceso se rechaza."),
], "Media", "CRUD técnico / Permisos")

# BLOQUE 16 — WhatsApp huésped
add_many("BLOQUE 16 — WhatsApp huésped", "Huésped", "WhatsApp huésped", [
    ("responder al saludo", "Envíe hola", "Recibe un único saludo y el menú principal con Ver opciones."),
    ("mostrar las nueve opciones", "Pulse Ver opciones", "Aparecen Alojamientos, Reservar ahora, Experiencias, Contacto, Clima, Preguntas frecuentes, Compartir experiencia, Mi reserva y Beneficios."),
    ("volver al menú con comandos globales", "Entre a una opción|Escriba menú, inicio y menu principal en pruebas separadas", "Cada comando vuelve al menú principal y cierra el paso anterior."),
    ("cancelar una conversación parcial", "Inicie Reserva|Escriba cancelar", "Cancela el paso en curso, no una reserva ya registrada, y vuelve al menú."),
    ("volver al paso anterior", "Entre a Alojamientos o Reserva|Escriba volver", "Vuelve de forma coherente sin conservar datos parciales incompatibles."),
    ("rechazar opción inexistente", "En el menú escriba 99 o texto desconocido", "Explica cómo usar botones o números 1 a 9."),
    ("listar alojamientos", "Seleccione Alojamientos", "Muestra los tipos activos con capacidad y precio."),
    ("mostrar detalle y galería por alojamiento", "Abra cada tipo activo uno por uno", "Cada tipo muestra su información y hasta el límite configurado de fotos propias."),
    ("usar botones del alojamiento", "Pulse Ver alojamientos, Reservar y Menú principal en recorridos separados", "Cada botón abre el destino correcto."),
    ("mostrar experiencias", "Seleccione Experiencias|Abra una actividad", "Muestra únicamente actividades activas con detalle y multimedia válida."),
    ("mostrar contacto", "Seleccione Contacto", "Muestra números públicos configurados y botón Menú principal."),
    ("mostrar clima o respaldo", "Seleccione Clima", "Si OpenWeather está configurado muestra pronóstico; si no, muestra indisponibilidad sin fallar el bot."),
    ("mostrar preguntas frecuentes", "Seleccione Preguntas frecuentes", "Muestra reglas reales, anticipo del 50 % y botones Reservar/Menú principal."),
    ("mostrar instrucciones para compartir experiencia", "Seleccione Compartir experiencia", "Explica el canal disponible sin prometer una integración inexistente."),
    ("consultar Mi reserva sin registros", "Desde un teléfono sin reserva seleccione Mi reserva", "Informa que no encontró reservas y ofrece agente o menú."),
    ("consultar Mi reserva pendiente", "Desde el teléfono de una solicitud seleccione Mi reserva", "Muestra código, estado y acciones permitidas según la etapa."),
    ("mostrar Beneficios", "Seleccione Beneficios", "Informa que promociones se consultan por canales oficiales y ofrece Contactar."),
    ("mantener orden de dos mensajes rápidos", "Envíe dos mensajes consecutivos válidos", "Las respuestas respetan el orden y no mezclan estados."),
    ("evitar respuesta duplicada", "Reenvíe técnicamente el mismo message_id en QA", "El huésped recibe una sola respuesta.", {"technical": True}),
], "Alta", "Conversación / Navegación / Multimedia")

# BLOQUE 17 — Reserva por WhatsApp
add_many("BLOQUE 17 — Reserva y aprobación WhatsApp", "Huésped", "WhatsApp huésped", [
    ("iniciar una reserva", "Seleccione Reservar ahora", "Solicita fechas futuras en un formato comprensible."),
    ("rechazar fechas incompletas", "Envíe una sola fecha o texto sin rango", "Explica el formato y mantiene el paso de fechas."),
    ("rechazar fechas pasadas", "Envíe un rango completamente pasado", "No crea solicitud y pide fechas futuras."),
    ("confirmar fechas válidas", "Envíe un rango futuro libre|Pulse Sí, confirmar", "Solicita nombre completo y conserva exactamente las fechas confirmadas."),
    ("corregir fechas con No", "Envíe fechas válidas|Pulse No", "Permite introducir otro rango y descarta el anterior."),
    ("validar nombre del huésped", "Envíe un carácter o solo números|Luego envíe QA Huésped 001", "Rechaza el nombre inválido y acepta el nombre completo."),
    ("validar cantidad de huéspedes", "Pruebe 0, texto y una cantidad sobre la capacidad|Luego envíe 2", "Rechaza valores inválidos y acepta una cantidad permitida."),
    ("mostrar resumen y total", "Complete fechas, nombre y personas", "Muestra alojamiento asignado, fechas, noches, personas y precio total."),
    ("rechazar condiciones", "En el resumen seleccione No acepto", "No crea reserva y vuelve a una salida segura."),
    ("crear solicitud al aceptar", "En el resumen pulse Acepto", "Crea un código VJ único y estado Espera autorización; todavía no solicita pago."),
    ("impedir doble solicitud por repetición", "Pulse Acepto dos veces o repita el mismo evento", "Solo existe una solicitud con un código."),
    ("impedir enviar comprobante antes de autorizar", "Envíe JPG antes de la aprobación", "Indica que debe esperar autorización y no adjunta el archivo."),
    ("recibir autorización e instrucciones", "Espere que el administrador autorice", "Recibe 50 %, importe, saldo, cuentas, notas y plazo de 24 horas."),
    ("aceptar comprobante JPG", "Después de autorizar, envíe JPG válido", "Confirma recepción y cambia a revisión final."),
    ("aceptar comprobante PNG", "En otra solicitud autorizada envíe PNG válido", "Confirma recepción."),
    ("aceptar comprobante PDF", "En otra solicitud autorizada envíe PDF válido", "Confirma recepción."),
    ("rechazar archivo inválido", "Después de autorizar envíe un formato no permitido", "No lo guarda y explica JPG, PNG o PDF."),
    ("rechazar archivo mayor de 5 MiB", "Envíe un archivo mayor al límite", "No lo guarda y muestra el límite sin bloquear la conversación."),
    ("recibir confirmación final", "Espere que el administrador confirme", "Recibe Reserva confirmada con código y próximos pasos."),
    ("recibir rechazo con motivo", "En una solicitud QA haga que el administrador rechace", "Recibe el motivo correcto una sola vez."),
    ("solicitar modificación o asistencia", "Desde Mi reserva seleccione Modificar reserva o Solicitar asistencia", "Los administradores reciben aviso y el huésped es informado de contacto privado."),
    ("informar política sin reembolso", "Desde Mi reserva solicite Cancelar reserva", "Informa la política de no reembolsos y deriva a atención; no promete devolución automática."),
], "Crítica", "Flujo completo / Archivos / Reglas")

# BLOQUE 18 — Administrador por WhatsApp
add_many("BLOQUE 18 — Administrador por WhatsApp", "Administrador WhatsApp", "WhatsApp administrador", [
    ("entrar al modo administrador", "Envíe /admin", "Recibe comandos y hasta cinco solicitudes pendientes."),
    ("rechazar /admin desde número no autorizado", "Desde un teléfono huésped envíe /admin", "No obtiene detalles administrativos ni botones de decisión."),
    ("consultar una solicitud", "Envíe /reserva VJ-XXXXXX", "Muestra únicamente la solicitud solicitada con huésped, cabaña, fechas, personas y estado."),
    ("consultar código inexistente", "Envíe /reserva VJ-999999", "Informa que no existe sin revelar otros datos."),
    ("autorizar pago con botón", "Abra una revisión previa|Pulse Autorizar pago", "La reserva pasa a esperando_pago y el huésped recibe instrucciones."),
    ("autorizar pago con comando", "Envíe /aprobar VJ-XXXXXX en una solicitud pendiente", "Ejecuta la autorización previa una sola vez."),
    ("rechazar con botón y motivo", "Pulse Rechazar|Escriba un motivo válido", "La solicitud queda rechazada y el huésped recibe el motivo."),
    ("rechazar con comando", "Envíe /rechazar VJ-XXXXXX MOTIVO QA", "La solicitud se rechaza y queda trazabilidad del administrador."),
    ("cancelar captura del motivo", "Pulse Rechazar|Escriba cancelar", "Cancela el rechazo y conserva el estado previo."),
    ("recibir segunda revisión", "Haga que el huésped envíe comprobante", "El administrador recibe VERIFICACIÓN DEL COMPROBANTE con Confirmar, Rechazar y Ver detalles."),
    ("confirmar con botón", "Abra el comprobante|Pulse Confirmar", "La reserva queda confirmada y el huésped es notificado o el mensaje se encola."),
    ("confirmar con comando", "Envíe /confirmar VJ-XXXXXX", "Confirma únicamente si el comprobante existe y el estado corresponde."),
    ("impedir confirmar sin archivo", "Use /confirmar sobre una reserva esperando_pago", "Responde RECEIPT_REQUIRED o mensaje equivalente y no confirma."),
    ("impedir una acción en estado inválido", "Repita /aprobar en una reserva ya confirmada", "No repite la acción ni envía mensajes duplicados."),
    ("volver al modo huésped", "Envíe /cliente|Luego envíe menú", "Recibe la experiencia normal de huésped y no acciones administrativas automáticas."),
], "Crítica", "WhatsApp / Permisos / Aprobación")

# BLOQUE 19 — Integraciones y seguridad técnica
add_many("BLOQUE 19 — Webhook e integraciones", "Superadministrador + técnico QA", "Webhook Meta", [
    ("verificar el webhook con token correcto", "En QA envíe la verificación GET con hub.mode, challenge y verify token válidos", "Responde 200 y devuelve exactamente el challenge.", {"technical": True}),
    ("rechazar verify token incorrecto", "Repita con token incorrecto", "Responde 403 sin revelar el token correcto.", {"technical": True}),
    ("aceptar POST con firma válida", "Envíe un evento QA firmado con META_APP_SECRET", "Responde 200 y procesa una sola vez.", {"technical": True}),
    ("rechazar POST sin firma", "Envíe el mismo cuerpo sin X-Hub-Signature-256", "Responde 401 y no procesa mensaje.", {"technical": True}),
    ("rechazar firma incorrecta", "Modifique el cuerpo después de firmarlo", "Responde 401 y no guarda el evento como procesado.", {"technical": True}),
    ("deduplicar message_id", "Envíe dos veces el mismo evento firmado", "La segunda entrega no genera respuesta adicional.", {"technical": True}),
    ("reintentar error temporal 429", "Simule respuesta 429 de Meta en QA", "El mensaje queda pending y aumenta next_attempt_at.", {"technical": True}),
    ("reintentar error 5xx o timeout", "Simule 500 y timeout de Meta", "El mensaje se conserva para reintento con error sanitizado.", {"technical": True}),
    ("no reintentar error permanente 400", "Simule 400 no recuperable", "Se marca fallido sin bucle infinito.", {"technical": True}),
    ("usar respaldo de texto si fallan botones", "Simule rechazo de mensaje interactivo", "Envía texto numerado equivalente una sola vez.", {"technical": True}),
    ("manejar clima no disponible", "Desactive OPENWEATHER_API_KEY en QA|Seleccione Clima", "El bot muestra respaldo y el resto del menú continúa."),
    ("manejar Dropbox no configurado", "Deje BACKUP_DROPBOX_ACCESS_TOKEN vacío en QA|Abra Copias", "Muestra Copia externa Pendiente sin afectar copias locales."),
], "Crítica", "Integración / Resiliencia / Seguridad")
add_many("BLOQUE 19 — Webhook e integraciones", "Superadministrador + técnico QA", "Seguridad", [
    ("responder health y readiness", "Consulte /health y /ready", "Health responde ok; Ready confirma database ok y estado de WhatsApp sin secretos.", {"technical": True}),
    ("aplicar CORS permitido", "Desde el origen configurado haga una solicitud preflight", "Responde con el origen permitido y no con comodín.", {"technical": True}),
    ("rechazar un origen CORS no permitido", "Repita desde un origen QA no autorizado", "No concede acceso CORS.", {"technical": True}),
    ("proteger rutas administrativas sin JWT", "Solicite /admin/reservations sin Authorization", "Responde 401 y no devuelve datos.", {"technical": True}),
    ("rechazar JWT inválido o expirado", "Use un token alterado o expirado", "Responde 401/403 y no devuelve datos.", {"technical": True}),
    ("invalidar sesión al desactivar cuenta", "Desactive una cuenta QA con sesión abierta|Repita una consulta", "La sesión se rechaza inmediatamente.", {"technical": True}),
    ("no registrar secretos", "Revise logs QA después de login, webhook y error Meta", "No aparecen contraseñas, JWT completos, access tokens, PIN ni app secret.", {"technical": True}),
    ("bloquear traversal en descarga", "Solicite un nombre de backup con ../", "Responde 404 y no descarga archivos fuera del directorio.", {"technical": True}),
    ("proteger comprobantes contra tipos no permitidos", "Intente subir un ejecutable renombrado", "No se acepta como comprobante válido.", {"technical": True}),
], "Crítica", "Seguridad / API")

# BLOQUE 20 — Matriz negativa de permisos
add_many("BLOQUE 20 — Permisos entre roles", "Administrador", "Seguridad", [
    ("ocultar Administradores al rol admin", "Inicie sesión como Admin|Revise menú|Abra /admin-users directamente", "El menú no aparece, la URL no abre y la API responde 403."),
    ("ocultar Configuración al rol admin", "Revise menú|Abra /settings directamente", "No puede ver ni cambiar pagos o números WhatsApp."),
    ("ocultar Copias al rol admin", "Revise menú|Abra /backups directamente", "No puede listar, crear ni descargar copias."),
    ("ocultar Mensajes pendientes al rol admin", "Revise menú|Abra /notifications directamente", "No puede ver destinatarios ni reintentar mensajes."),
    ("ocultar Auditoría al rol admin", "Revise menú|Abra /audit-logs directamente", "No puede consultar la bitácora."),
    ("ocultar Estados de conversación al rol admin", "Abra /conversation-states directamente", "No puede leer ni modificar conversaciones."),
], "Crítica", "Permisos negativos")

# BLOQUE 21 — Interfaz y casos límite
add_many("BLOQUE 21 — Interfaz y casos límite", "Todos los roles", "Interfaz", [
    ("mostrar el panel correctamente en escritorio", "Use 1440×900|Recorra todos los menús permitidos", "No hay contenido cortado, superpuesto ni botones fuera de pantalla."),
    ("mostrar el panel correctamente en portátil", "Use 1366×768|Recorra formularios y tablas", "Los controles siguen accesibles y las tablas permiten desplazamiento cuando corresponde."),
    ("mostrar el panel en teléfono", "Use aproximadamente 390×844|Abra menú lateral, formularios y modales", "El menú móvil abre/cierra y no impide guardar o cancelar."),
    ("mostrar textos legibles y no técnicos", "Revise mensajes vacíos, errores y confirmaciones", "No muestra stack traces, SQL, undefined, NaN ni objetos JSON al operador normal."),
    ("manejar doble clic", "En QA haga doble clic en Guardar o Confirmar", "El botón se deshabilita o el backend evita duplicados."),
    ("manejar actualización durante una operación", "Inicie una carga no destructiva y actualice", "Al volver, el estado es consistente; no crea registros parciales duplicados."),
    ("manejar pérdida temporal de conexión", "Desconecte red en QA|Pulse Actualizar|Reconecte y repita", "Muestra error claro y se recupera sin recargar datos falsos."),
    ("manejar caracteres especiales", "Use QA O'Connor & Familia en un nombre permitido|Guarde", "El texto se conserva sin romper consultas ni interfaz."),
    ("manejar texto extremadamente largo", "Pegue texto superior al límite en descripción o nombre|Guarde", "Valida o recorta de forma visible; no rompe tablas ni servidor."),
    ("manejar varias pestañas", "Edite el mismo registro QA en dos pestañas de forma controlada", "El último resultado queda consistente y no corrompe relaciones; cualquier conflicto es visible."),
], "Media", "Interfaz / Casos límite")

# BLOQUE 22 — Flujos cruzados, regresión y salida
add_many("BLOQUE 22 — Pruebas cruzadas", "Todos los roles", "Regresión", [
    ("completar una reserva de extremo a extremo", "Huésped crea solicitud|Admin WhatsApp autoriza|Huésped paga y adjunta|Admin confirma|Admin panel revisa calendario y auditoría", "Todos los canales muestran el mismo código, estado, fechas, importe y responsables."),
    ("rechazar una solicitud de extremo a extremo", "Huésped crea solicitud|Administrador rechaza con motivo|Abra panel y Mi reserva", "El motivo y estado rechazado coinciden en todos los canales."),
    ("recuperar un aviso temporalmente fallido", "Simule error temporal Meta|Confirme que aparece en Mensajes pendientes|Restablezca servicio|Reintente", "El aviso pasa a enviado sin repetir la transición de la reserva."),
    ("reflejar cambios de contenido en WhatsApp", "Edite un tipo o actividad QA|Abra vista previa|Consulte desde huésped", "Panel, vista previa y WhatsApp muestran el mismo contenido activo."),
    ("reflejar una reserva en reportes", "Cree y confirme una reserva QA|Abra Dashboard, Calendario y Reportes", "Los tres módulos reflejan el mismo registro sin doble conteo."),
    ("registrar acciones superiores en auditoría", "Cree admin QA, cambie pagos, cree backup y reintente mensaje|Abra Auditoría", "Cada acción aparece con usuario, ruta y resultado."),
    ("persistir después de reinicio controlado QA", "Anote datos QA|Solicite al técnico reiniciar solo el ambiente QA|Vuelva a consultar", "Usuarios, reservas, estados, auditoría y cola persisten.", {"technical": True}),
    ("ejecutar la regresión automatizada", "El técnico ejecuta pnpm run validate en backend y pnpm validate en frontend con VITE_API_URL QA", "Todas las suites, compilación y auditorías terminan con código 0.", {"technical": True}),
], "Crítica", "Flujo cruzado / Regresión")
add_many("BLOQUE 23 — Validación final de producción", "Superadministrador + técnico QA", "Producción", [
    ("confirmar datos bancarios reales", "Abra Configuración de producción sin cambiar datos|Verifique que no exista PRUEBA o NO PAGAR", "Todas las cuentas y titulares fueron confirmados por el propietario."),
    ("confirmar credenciales individuales", "Verifique que cada operador tenga cuenta propia y haya cambiado la contraseña temporal", "No se comparte la contraseña de demostración."),
    ("confirmar token permanente y app Meta activa", "En Meta verifique usuario del sistema, token permanente, número definitivo, modo activo y webhook messages", "La integración no depende de token temporal y el número tiene estado conectado."),
    ("confirmar plantilla administrativa de Meta", "Verifique una plantilla Utility aprobada y configurada para alertas fuera de 24 horas", "Un administrador recibe una alerta aun sin conversación reciente."),
    ("confirmar copia externa", "Verifique Dropbox activo o una descarga externa reciente", "Existe al menos una copia recuperable fuera de Railway."),
    ("confirmar salud pública", "Abra /health y /ready", "Ambos responden correctamente sin datos sensibles."),
    ("ejecutar humo no destructivo en producción", "Inicie sesión|Abra Dashboard, Reservas, Reportes y Copias|Abra menú WhatsApp con un teléfono QA", "Las lecturas y navegación funcionan sin modificar clientes ni reservas reales."),
    ("tomar la decisión final de salida", "Complete la tabla de resultados|Calcule porcentaje|Revise bugs críticos y altos|Marque SÍ, NO o LISTO CON OBSERVACIONES", "Solo se aprueba si se cumplen los criterios de salida definidos al final del manual."),
], "Crítica", "Aceptación / Producción")


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for name, size, color in (("Heading 1", 16, DARK_BLUE), ("Heading 2", 12.5, BLUE), ("Heading 3", 10.5, DARK_BLUE)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(9)
        style.paragraph_format.space_after = Pt(4)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    set_run_font(p.add_run("VILLAS JULIE"), 15, True, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Manual Maestro de Pruebas QA/UAT"), 25, True, DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Procedimiento ejecutable, checklist y decisión de producción"), 13, False, GREEN)
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    set_table_widths(table, [2600, 6760])
    rows = [
        ("Versión", "1.0 — 27 de agosto de 2026"),
        ("Sistema", "Panel administrativo, API, SQLite y WhatsApp Business Cloud de Meta"),
        ("Ambiente", "QA separado. Producción solo para el bloque final no destructivo"),
        ("Estado inicial", "Todas las pruebas manuales: ⬜ No ejecutada"),
        ("Propietario", "Carlos Velasquez — Villas Julie"),
    ]
    for row, values in zip(table.rows, rows):
        for idx, value in enumerate(values):
            row.cells[idx].text = value
            if idx == 0:
                set_cell_shading(row.cells[idx], DARK_BLUE)
                for run in row.cells[idx].paragraphs[0].runs:
                    set_run_font(run, 9.5, True, WHITE)
            else:
                for run in row.cells[idx].paragraphs[0].runs:
                    set_run_font(run, 9.5)
    add_callout(doc, "REGLA PRINCIPAL", "No ejecute pruebas de creación, edición, eliminación, fallos simulados o archivos inválidos en producción. Una prueba manual permanece No ejecutada hasta que una persona complete el procedimiento y guarde evidencia.", YELLOW)
    doc.add_page_break()


def add_intro(doc):
    doc.add_heading("1. Cómo usar este manual", level=1)
    for text in [
        "Ejecute las pruebas en orden desde PRUEBA 001 hasta la última. No salte un requisito previo porque los bloques posteriores reutilizan datos creados antes.",
        "Después de cada procedimiento marque Aprobada, Fallida, Bloqueada o Pendiente; escriba el resultado obtenido y la ubicación de la evidencia.",
        "Si una prueba falla, cree un BUG con el formato del capítulo de incidencias. Después de corregirlo ejecute la prueba original y el conjunto de regresión indicado.",
        "Las pruebas técnicas están claramente identificadas. El propietario puede seguirlas con apoyo de una persona técnica sin necesidad de conocer el código.",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "Interpretación", "La existencia de código o una prueba automatizada no aprueba una prueba UAT. La aprobación requiere ejecución manual, resultado esperado y evidencia.")

    doc.add_heading("2. Arquitectura y roles reales encontrados", level=1)
    architecture = [
        ("Panel", "React + Vite en Cloudflare Pages", "Operación administrativa"),
        ("API", "Node.js + Express en Railway", "Reglas, seguridad, reportes y WhatsApp"),
        ("Base de datos", "SQLite persistente con WAL", "Datos, estados, auditoría y cola"),
        ("Mensajería", "WhatsApp Business Cloud API de Meta", "Atención de huéspedes y decisiones administrativas"),
        ("Archivos", "Volumen Railway", "Comprobantes y copias"),
        ("Opcionales", "Dropbox y OpenWeather", "Copia externa y clima"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text, table.rows[0].cells[1].text, table.rows[0].cells[2].text = "Componente", "Tecnología", "Uso"
    for row in architecture:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_widths(table, [1800, 3000, 4560])
    style_table(table)
    doc.add_paragraph()
    roles = [
        ("Superadministrador", "Panel", "Todo lo operativo más cuentas, pagos, números WhatsApp, auditoría, cola, estados y copias"),
        ("Administrador", "Panel", "Dashboard, reservas, calendario, huéspedes, cabañas, tipos, actividades y reportes"),
        ("Administrador WhatsApp", "Teléfono autorizado", "Autorizar pago, consultar, rechazar y confirmar reservas"),
        ("Huésped", "WhatsApp", "Consultar, reservar, enviar comprobante autorizado y recibir decisiones"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text, table.rows[0].cells[1].text, table.rows[0].cells[2].text = "Rol real", "Canal", "Alcance"
    for row in roles:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_widths(table, [2200, 1900, 5260])
    style_table(table)


def add_preparation(doc):
    doc.add_heading("3. Preparación del ambiente", level=1)
    add_label_paragraph(doc, "Panel producción", "https://admin-frontend-cgh.pages.dev/")
    add_label_paragraph(doc, "API producción", "https://vj-production-49d7.up.railway.app")
    add_label_paragraph(doc, "Ambiente QA", "PENDIENTE DE CREAR o confirmar; debe usar base y almacenamiento separados")
    add_label_paragraph(doc, "Navegadores", "Chrome o Edge actual; ventana privada; resolución 1440×900 y 1366×768")
    add_label_paragraph(doc, "Teléfonos", "Un administrador WhatsApp QA y un huésped QA autorizados por Meta")
    doc.add_heading("3.1 Credenciales y teléfonos", level=2)
    creds = [
        ("Superadministrador", "admin (confirmar)", "PENDIENTE DE CONFIRMAR/CAMBIAR", "Existente", "No escribir la contraseña en evidencias"),
        ("Administrador", "PENDIENTE DE CREAR", "PENDIENTE DE CREAR", "Pendiente", "Cuenta QA sin privilegios superiores"),
        ("Administrador WhatsApp", "Carlos / Gregorio o número QA", "No usa PIN del sistema", "Confirmar", "Debe enviar /admin"),
        ("Huésped", "PENDIENTE DE CREAR", "No aplica", "Pendiente", "Número autorizado por Meta"),
    ]
    table = doc.add_table(rows=1, cols=5)
    for i, value in enumerate(("Rol", "Usuario de prueba", "Contraseña/PIN", "Estado", "Observaciones")):
        table.rows[0].cells[i].text = value
    for row in creds:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_widths(table, [1500, 1900, 1900, 1200, 2860])
    style_table(table)
    doc.add_heading("3.2 Información que necesito del propietario", level=2)
    for item in [
        "URL o confirmación del ambiente QA separado.",
        "Contraseñas QA definidas fuera de este documento.",
        "Número huésped QA y número administrador QA autorizados por Meta.",
        "Confirmación de que el token Meta es permanente y la aplicación está activa.",
        "Cuentas bancarias reales solamente para la validación final; no copiarlas en evidencias.",
        "Plantilla Utility aprobada para alertas administrativas fuera de 24 horas.",
        "Responsable y ubicación para guardar las evidencias y copias externas.",
    ]:
        add_bullet(doc, item, checked=True)


def add_permissions(doc):
    doc.add_heading("4. Matriz de permisos", level=1)
    rows = [
        ("Dashboard, reservas, calendario, huéspedes", "✅", "✅", "❌", "❌"),
        ("Cabañas, tipos, actividades y reportes", "✅", "✅", "❌", "Consulta indirecta"),
        ("Administradores del panel", "✅", "❌", "❌", "❌"),
        ("Pagos y administradores WhatsApp", "✅", "❌", "❌", "❌"),
        ("Copias, cola, auditoría y estados", "✅", "❌", "❌", "❌"),
        ("Autorizar/confirmar por WhatsApp", "❌", "❌", "✅", "❌"),
        ("Menú y reserva por WhatsApp", "Como huésped", "Como huésped", "Con /cliente", "✅"),
        ("Enviar comprobante autorizado", "Como huésped", "Como huésped", "Con /cliente", "✅"),
    ]
    table = doc.add_table(rows=1, cols=5)
    for i, value in enumerate(("Función", "Superadmin", "Admin", "Admin WhatsApp", "Huésped")):
        table.rows[0].cells[i].text = value
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_widths(table, [3760, 1400, 1200, 1500, 1500])
    style_table(table)
    add_callout(doc, "Verificación obligatoria", "Toda restricción importante se prueba en tres niveles: menú oculto, URL directa bloqueada y API rechazada. Ocultar un botón no se considera seguridad suficiente.")


def add_inventory(doc):
    doc.add_heading("5. Inventario maestro de funcionalidades y cobertura", level=1)
    p = doc.add_paragraph("Cada fila corresponde a una función o escenario encontrado y posee una prueba asociada. La tabla completa continúa en las páginas siguientes.")
    p.paragraph_format.space_after = Pt(5)
    table = doc.add_table(rows=1, cols=7)
    headers = ("ID", "Módulo", "Función", "Rol autorizado", "Acción", "Criticidad", "Prueba")
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
    for test in tests:
        cells = table.add_row().cells
        values = (
            f"F-{test['id']:03d}", test["module"], test["title"].capitalize(), test["role"],
            test["type"], test["criticality"], f"PRUEBA {test['id']:03d}",
        )
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_widths(table, [600, 1250, 2480, 1400, 1300, 1000, 1330])
    style_table(table)


def add_test(doc, test):
    doc.add_heading(f"PRUEBA {test['id']:03d} — {test['title'].capitalize()}", level=2)
    meta = doc.add_table(rows=2, cols=4)
    labels = (("Rol", test["role"], "Módulo", test["module"]), ("Criticidad", test["criticality"], "Tipo", test["type"]))
    for r_idx, row in enumerate(labels):
        for c_idx, value in enumerate(row):
            meta.rows[r_idx].cells[c_idx].text = value
            if c_idx % 2 == 0:
                set_cell_shading(meta.rows[r_idx].cells[c_idx], DARK_BLUE)
                for run in meta.rows[r_idx].cells[c_idx].paragraphs[0].runs:
                    set_run_font(run, 9, True, WHITE)
            else:
                for run in meta.rows[r_idx].cells[c_idx].paragraphs[0].runs:
                    set_run_font(run, 9)
    set_table_widths(meta, [1200, 3300, 1200, 3660])
    if test["technical"]:
        add_callout(doc, "APOYO TÉCNICO", "Esta prueba no debe ejecutarse sin una persona técnica y nunca debe simular fallos contra producción.", YELLOW)
    doc.add_heading("Objetivo", level=3)
    doc.add_paragraph(test["objective"])
    doc.add_heading("Requisitos previos", level=3)
    for item in test["requirements"]:
        add_bullet(doc, item, checked=True)
    doc.add_heading("Datos que utilizaremos", level=3)
    data = doc.add_table(rows=1, cols=2)
    data.rows[0].cells[0].text, data.rows[0].cells[1].text = "Campo", "Valor de prueba"
    for key, value in test["data"]:
        cells = data.add_row().cells
        cells[0].text, cells[1].text = key, value
    set_table_widths(data, [2500, 6860])
    style_table(data)
    doc.add_heading("Procedimiento", level=3)
    for index, step in enumerate(test["steps"], 1):
        add_bullet(doc, f"{index}. {step}", checked=True)
    doc.add_heading("Resultado esperado", level=3)
    doc.add_paragraph(test["expected"])
    doc.add_heading("Cómo verificarlo", level=3)
    for item in test["verify"]:
        add_bullet(doc, item, checked=True)
    doc.add_heading("El resultado indica error si", level=3)
    for item in test["failure"]:
        add_bullet(doc, item)
    doc.add_heading("Evidencia que debe guardar", level=3)
    doc.add_paragraph(test["evidence"])
    result = doc.add_table(rows=5, cols=2)
    rows = [
        ("Estado", "☐ APROBADA   ☐ FALLIDA   ☐ BLOQUEADA   ☐ PENDIENTE"),
        ("Resultado obtenido", "____________________________________________________________"),
        ("Error / BUG", "____________________________________________________________"),
        ("Evidencia", "____________________________________________________________"),
        ("Fecha / Probado por", "____________________________________________________________"),
    ]
    for row, values in zip(result.rows, rows):
        row.cells[0].text, row.cells[1].text = values
        set_cell_shading(row.cells[0], GRAY)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, 9, True, DARK_BLUE)
    set_table_widths(result, [2300, 7060])
    doc.add_paragraph().paragraph_format.space_after = Pt(5)


def add_bug_and_regression(doc):
    doc.add_heading("Registro de errores", level=1)
    doc.add_heading("BUG-001 — Nombre del problema", level=2)
    for label in ("Prueba relacionada", "Módulo", "Rol", "Severidad: Crítica / Alta / Media / Baja", "Descripción", "Pasos para reproducir", "Resultado esperado", "Resultado obtenido", "Evidencia"):
        add_label_paragraph(doc, label, "____________________________________________________________")
    add_label_paragraph(doc, "Estado", "☐ Nuevo  ☐ En revisión  ☐ Corregido  ☐ Pendiente de volver a probar  ☐ Cerrado")
    doc.add_heading("Regresión mínima después de una corrección", level=1)
    regression = [
        ("Autenticación/roles", "PRUEBAS 009–022 y todas las pruebas de permisos"),
        ("Reservas/precios", "Listado, creación, conflicto, autorización, comprobante, confirmación, calendario, dashboard y reportes"),
        ("WhatsApp", "Menú, navegación, galería, reserva, webhook, deduplicación, cola y administración"),
        ("Archivos", "Fotos de cabaña, multimedia, comprobantes y copias"),
        ("Base de datos", "Persistencia, relaciones, no superposición, auditoría y copia"),
        ("Frontend", "Prueba modificada, navegación, permisos, actualización y dispositivo móvil"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Área corregida", "Regresión obligatoria"
    for row in regression:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = row
    set_table_widths(table, [2500, 6860])
    style_table(table)


def add_master_results(doc):
    doc.add_heading("Checklist maestra", level=1)
    groups = []
    for test in tests:
        if test["group"] not in groups:
            groups.append(test["group"])
    for group in groups:
        count = sum(1 for t in tests if t["group"] == group)
        add_bullet(doc, f"{group}: {count} pruebas completadas y con evidencia.", checked=True)
    for item in [
        "Ambiente QA separado y respaldado.",
        "Todos los roles y credenciales preparados.",
        "Pruebas positivas, negativas y permisos completadas.",
        "Bugs críticos y altos corregidos y revalidados.",
        "Regresión automatizada y manual completada.",
        "Copia externa y datos reales confirmados.",
    ]:
        add_bullet(doc, item, checked=True)

    doc.add_heading("Tabla maestra de resultados", level=1)
    table = doc.add_table(rows=1, cols=7)
    for i, value in enumerate(("ID", "Prueba", "Rol", "Módulo", "Criticidad", "Estado", "Bug")):
        table.rows[0].cells[i].text = value
    for test in tests:
        cells = table.add_row().cells
        values = (f"{test['id']:03d}", test["title"].capitalize(), test["role"], test["module"], test["criticality"], "⬜ No ejecutada", "—")
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_widths(table, [550, 2770, 1350, 1250, 900, 1500, 1040])
    style_table(table)
    doc.add_paragraph()
    totals = doc.add_table(rows=6, cols=2)
    summary = [
        ("Total de pruebas", str(len(tests))),
        ("Aprobadas", "_____"),
        ("Fallidas", "_____"),
        ("Bloqueadas", "_____"),
        ("Pendientes", "_____"),
        ("Porcentaje aprobado", "Aprobadas ÷ Total × 100 = ______ %"),
    ]
    for row, values in zip(totals.rows, summary):
        row.cells[0].text, row.cells[1].text = values
        set_cell_shading(row.cells[0], DARK_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, 9.5, True, WHITE)
    set_table_widths(totals, [3000, 6360])


def add_release_criteria(doc):
    doc.add_heading("Criterios para producción", level=1)
    severity = [
        ("CRÍTICO", "Seguridad, pérdida/corrupción de datos, acceso, reserva, autorización, pago, comprobante o confirmación esencial.", "Debe ser 0 abierto"),
        ("ALTO", "Función importante inutilizable o resultado incorrecto sin alternativa segura.", "Debe ser 0 abierto"),
        ("MEDIO", "Problema con alternativa documentada y riesgo controlado.", "Solo con aceptación del propietario"),
        ("BAJO", "Detalle visual o menor sin impacto operativo.", "Puede quedar planificado"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text, table.rows[0].cells[1].text, table.rows[0].cells[2].text = "Nivel", "Definición", "Salida"
    for row in severity:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_widths(table, [1300, 5860, 2200])
    style_table(table)
    doc.add_heading("Puerta de aprobación", level=2)
    for item in [
        "100 % de pruebas críticas ejecutadas y aprobadas.",
        "100 % de pruebas altas ejecutadas y aprobadas.",
        "Al menos 95 % del total aprobado; ninguna prueba crítica bloqueada.",
        "0 bugs críticos y 0 bugs altos abiertos.",
        "Prueba completa real de WhatsApp aprobada con evidencia.",
        "Cuentas reales verificadas, contraseñas individuales, token Meta permanente y plantilla administrativa aprobada.",
        "Copia externa recuperable y health/ready correctos.",
    ]:
        add_bullet(doc, item, checked=True)
    add_callout(doc, "¿ESTÁ LISTO PARA PRODUCCIÓN?", "☐ SÍ     ☐ NO     ☐ LISTO CON OBSERVACIONES\n\nPorcentaje aprobado: ______ %\nBugs críticos: _____  Altos: _____  Medios: _____  Bajos: _____\nRiesgos aceptados: _________________________________________________\nRecomendación final: _______________________________________________", LIGHT_GREEN)


def build():
    doc = Document(REFERENCE)
    clear_body_keep_section(doc)
    setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    doc.core_properties.title = "Manual Maestro de Pruebas QA/UAT — Villas Julie"
    doc.core_properties.subject = "Plan ejecutable de pruebas y aprobación de producción"
    doc.core_properties.author = "Villas Julie"
    doc.core_properties.keywords = "QA, UAT, WhatsApp, reservas, producción"
    add_title(doc)
    add_intro(doc)
    add_preparation(doc)
    add_permissions(doc)
    add_inventory(doc)
    doc.add_page_break()
    doc.add_heading("6. Orden exacto de ejecución", level=1)
    groups = []
    for test in tests:
        if test["group"] not in groups:
            groups.append(test["group"])
    for idx, group in enumerate(groups, 1):
        ids = [t["id"] for t in tests if t["group"] == group]
        add_numbered(doc, f"{group}: PRUEBAS {min(ids):03d}–{max(ids):03d}")
    doc.add_page_break()
    current_group = None
    for test in tests:
        if test["group"] != current_group:
            if current_group is not None:
                doc.add_page_break()
            current_group = test["group"]
            doc.add_heading(current_group, level=1)
        add_test(doc, test)
    doc.add_page_break()
    add_bug_and_regression(doc)
    add_master_results(doc)
    add_release_criteria(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    validate_output()
    print(f"Manual creado: {OUTPUT}")
    print(f"Total de pruebas: {len(tests)}")


def validate_output():
    with ZipFile(OUTPUT) as archive:
        damaged = archive.testzip()
        if damaged:
            raise RuntimeError(f"Paquete DOCX dañado: {damaged}")
        required_parts = {
            "word/document.xml", "word/styles.xml", "word/numbering.xml",
            "word/header1.xml", "word/footer1.xml", "word/theme/theme1.xml",
        }
        missing = required_parts.difference(archive.namelist())
        if missing:
            raise RuntimeError(f"Partes DOCX ausentes: {sorted(missing)}")
    check = Document(OUTPUT)
    headings = [p.text for p in check.paragraphs if p.style.name == "Heading 2" and p.text.startswith("PRUEBA ")]
    expected = [f"PRUEBA {index:03d}" for index in range(1, len(tests) + 1)]
    actual = [heading.split(" — ", 1)[0] for heading in headings]
    if actual != expected:
        raise RuntimeError("Los IDs de prueba no son consecutivos o están duplicados")
    all_text = "\n".join(p.text for p in check.paragraphs)
    all_text += "\n" + "\n".join(cell.text for table in check.tables for row in table.rows for cell in row.cells)
    for required in ("Inventario maestro", "Matriz de permisos", "Tabla maestra de resultados", "¿ESTÁ LISTO PARA PRODUCCIÓN?", f"Total de pruebas\n{len(tests)}"):
        if required not in all_text:
            raise RuntimeError(f"Contenido obligatorio ausente: {required}")
    if OUTPUT.stat().st_size < 100_000:
        raise RuntimeError("El documento generado es inesperadamente pequeño")


if __name__ == "__main__":
    build()
